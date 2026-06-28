import streamlit as st
import re
import os
import fitz
import json
import time
from dotenv import load_dotenv
from google import genai
from docx import Document
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet

st.set_page_config(
    page_title="Resume Analyzer AI",
    page_icon="📄",
    layout="wide",
    initial_sidebar_state="expanded",
)

st.markdown(
    """
<style>

/* Hide Streamlit Header */

#MainMenu{
visibility:hidden;
}

footer{
visibility:hidden;
}

header{
visibility:hidden;
}

/* Cards */

.card{

background:#262730;

padding:20px;

border-radius:18px;

border:1px solid #444;

margin-bottom:15px;

box-shadow:0px 4px 18px rgba(0,0,0,.3);

}

/* Score */

.score{

font-size:60px;

font-weight:bold;

text-align:center;

}

.subtitle{

font-size:18px;

text-align:center;

opacity:.85;

}

</style>
""",
    unsafe_allow_html=True,
)


def card(title, content):

    st.markdown(
        f"""
<div style="
background:#262730;
padding:20px;
border-radius:15px;
border:1px solid #444;
margin-bottom:20px;
box-shadow:0px 3px 15px rgba(0,0,0,.25);
height:320px;
overflow-y:auto;
">

<h3>{title}</h3>

{content}

</div>
""",
        unsafe_allow_html=True,
    )


def scroll_card(title, content):

    st.markdown(
        f"""
<div style="
background:#262730;
padding:20px;
border-radius:15px;
border:1px solid #444;
margin-bottom:20px;
box-shadow:0px 3px 15px rgba(0,0,0,.25);
height:320px;
overflow-y:auto;
">

<h3>{title}</h3>

{content}

</div>
""",
        unsafe_allow_html=True,
    )


def badge_card(title, items):

    badges = ""

    for item in items:
        badges += f"""
<span style="
display:inline-block;
background:#4F46E5;
color:white;
padding:8px 14px;
border-radius:20px;
margin:6px;
font-size:15px;
font-weight:500;
">
{item}
</span>
"""

    scroll_card(title, badges)


load_dotenv()

api_key = os.getenv("GEMINI_API_KEY")

client = genai.Client(api_key=api_key)


st.markdown(
    """
<style>

.hero{
background:linear-gradient(90deg,#4F46E5,#7C3AED);
padding:35px;
border-radius:20px;
text-align:center;
color:white;
margin-bottom:25px;
box-shadow:0px 6px 20px rgba(0,0,0,.25);
}

.hero h1{
font-size:48px;
margin-bottom:8px;
}

.hero p{
font-size:20px;
opacity:.95;
}

</style>

<div class="hero">

<h1>🤖 Resume Analyzer AI</h1>

<p>
Analyze your resume with AI, calculate an ATS score, identify missing skills, and receive professional career recommendations.
</p>

</div>
""",
    unsafe_allow_html=True,
)

with st.sidebar:

    st.markdown("## 📌 Resume Analyzer AI")

    st.markdown("---")

    st.info("""
🚀 **Version:** 1.0

🤖 **Model:** Gemini 2.5 Flash

📄 **Supported:** PDF
""")

    st.write("""
Welcome to **Resume Analyzer AI**.

This application analyzes resumes using **Google Gemini AI**
and provides professional ATS-based feedback.
""")

    st.divider()

    st.subheader("⚙️ Technologies")

    st.markdown("""
- Google Gemini 2.5 Flash
- Streamlit
- PyMuPDF
- Python
""")

    st.divider()

    st.subheader("👨‍💻 Developer")

    st.success("👨‍💻 Vansh Garg")

    st.caption("AI & Machine Learning Developer")

left, right = st.columns([2, 1], gap="large")

with left:

    st.markdown("## 📄 Upload Resume")

    st.write("Upload your resume in PDF format to receive an AI-powered ATS analysis.")

    st.info("📂 Drag & Drop your resume here or click below to browse.")

    uploaded_file = st.file_uploader(
        "",
        type=["pdf"],
        label_visibility="collapsed",
    )

    if uploaded_file is not None:
        st.success("✅ Resume uploaded successfully!")
        st.write(f"**📄 File Name:** {uploaded_file.name}")
        st.write(f"**📦 File Size:** {round(uploaded_file.size/1024,2)} KB")

    analyze = st.button(
        "🚀 Analyze Resume",
        use_container_width=True,
    )

with right:

    st.markdown("## 💡 Tips")

    st.info("""
### Before Uploading

✅ Keep your resume to **1 page**

✅ Use ATS-friendly formatting

✅ Add measurable achievements

✅ Mention  projects

✅ Include GitHub & LinkedIn

✅ Use action verbs

✅ Save as PDF
""")

if analyze:

    if uploaded_file is None:

        st.warning("⚠️ Please upload a PDF.")

    else:
        with fitz.open(
            stream=uploaded_file.read(),
            filetype="pdf",
        ) as doc:

            page_count = doc.page_count

            text = ""
            for page in doc:
                text += page.get_text()

        word_count = len(text.split())

        if not text.strip():
            st.error("❌ No readable text found in the uploaded PDF.")
            st.stop()
        prompt = f"""
You are an expert ATS recruiter with over 15 years of experience hiring AI/ML engineers.

Analyze the following resume.

For certifications:

- Recommend 3 to 5 certifications suitable for the candidate.
- Even if the resume is excellent, always suggest certifications that can improve career growth.
- Never return an empty certifications list.

Never include markdown.
Never include explanations.
Never wrap JSON inside ```json.
Every list must contain at least one item.
ATS Score must be an integer from 0 to 100.

Return exactly in this structure:

{{
    "ats_score": 0,
    "summary": "",
    "technical_skills": [],
    "soft_skills": [],
    "missing_skills": [],
    "job_roles": [],
    "strengths": [],
    "weaknesses": [],
    "suggestions": [],
    "certifications": [],
    "final_verdict": ""
}}

Resume:

{text}
"""
        with st.spinner("🚀 AI is analyzing your resume..."):
            progress = st.progress(0)
            for i in range(70):
                progress.progress(i + 1)
                time.sleep(0.02)
            try:
                response = client.models.generate_content(
                    model="gemini-2.5-flash-lite",
                    contents=prompt,
                )
                progress.progress(100)

                json_text = response.text.strip()
                if json_text.startswith("```json"):
                    json_text = json_text.replace("```json", "")
                    json_text = json_text.replace("```", "")
                    json_text = json_text.strip()

                data = json.loads(json_text)

                ats_score = data["ats_score"]

                if ats_score >= 80:
                    score_color = "#00E676"

                elif ats_score >= 60:
                    score_color = "#FFC107"

                else:
                    score_color = "#F44336"

                st.success("✅ Resume analyzed successfully!")

                st.markdown("---")

                c1, c2, c3, c4, c5 = st.columns(5)
                with c1:
                    st.metric("⭐ ATS Score", f"{ats_score}/100")
                with c2:
                    st.metric("📄 Word Count", word_count)
                with c3:
                    st.metric("📑 Pages", page_count)
                with c4:
                    st.metric("💻 Skills Found", len(data["technical_skills"]))
                with c5:
                    st.metric("⚠ Missing Skills", len(data["missing_skills"]))
                    st.progress(ats_score / 100)
                    if ats_score >= 80:
                        st.success("🟢 Excellent ATS Compatibility")
                    elif ats_score >= 60:
                        st.warning("🟡 Average ATS Compatibility")
                    else:
                        st.error("🔴 Needs Improvement")

                st.markdown("---")

                status1, status2 = st.columns(2)
                with status1:
                    if ats_score >= 80:
                        st.success("🟢 Excellent ATS Compatibility")
                    elif ats_score >= 60:
                        st.warning("🟡 Average ATS Compatibility")
                    else:
                        st.error("🔴 Poor ATS Compatibility")

                with status2:
                    if ats_score >= 80:
                        level = "Strong"
                    elif ats_score >= 60:
                        level = "Average"
                    else:
                        level = "Needs Improvement"
                    st.metric(
                        "🎯 Resume Level",
                        level,
                    )

                st.markdown("---")

                st.subheader("📊 Resume Analysis")

                scroll_card(
                    "📝 Resume Summary",
                    data["summary"],
                )

                st.markdown("---")

                col1, col2 = st.columns(2)

                with col1:
                    with st.expander(
                        "💻 Technical Skills",
                        expanded=True,
                    ):
                        badge_card(
                            "",
                            data["technical_skills"],
                        )

                with col2:
                    with st.expander(
                        "🤝 Soft Skills",
                        expanded=True,
                    ):
                        badge_card(
                            "",
                            data["soft_skills"],
                        )

                col3, col4 = st.columns(2)

                with col3:
                    with st.expander(
                        "⚠ Missing Skills",
                        expanded=True,
                    ):
                        badge_card(
                            "",
                            data["missing_skills"],
                        )
                with col4:
                    with st.expander(
                        "💼 Best Job Roles",
                        expanded=True,
                    ):
                        badge_card(
                            "",
                            data["job_roles"],
                        )

                col5, col6 = st.columns(2)

                with col5:
                    with st.expander(
                        "💪 Strengths",
                        expanded=True,
                    ):
                        badge_card(
                            "",
                            data["strengths"],
                        )
                with col6:
                    with st.expander(
                        "📉 Weaknesses",
                        expanded=True,
                    ):
                        badge_card(
                            "",
                            data["weaknesses"],
                        )

                suggestions = "<ul>"
                for s in data["suggestions"]:
                    suggestions += f"<li>{s}</li>"
                suggestions += "</ul>"

                card(
                    "💡 Improvement Suggestions",
                    suggestions,
                )

                with st.expander(
                    "🎓 Certifications",
                    expanded=False,
                ):
                    if data["certifications"]:
                        badge_card(
                            "",
                            data["certifications"],
                        )
                    else:
                        st.info("No certifications recommended.")

                st.subheader("✅ Final Verdict")
                if ats_score >= 80:
                    st.success(data["final_verdict"])
                elif ats_score >= 60:
                    st.warning(data["final_verdict"])
                else:
                    st.error(data["final_verdict"])
                report = f"""
==============================
Resume Analysis Report
==============================

ATS Score:
{ats_score}/100

--------------------------------

Summary:
{data["summary"]}

--------------------------------

Technical Skills:
{", ".join(data["technical_skills"])}

--------------------------------

Soft Skills:
{", ".join(data["soft_skills"])}

--------------------------------

Missing Skills:
{", ".join(data["missing_skills"])}

--------------------------------

Best Job Roles:
{", ".join(data["job_roles"])}

--------------------------------

Strengths:
{", ".join(data["strengths"])}

--------------------------------

Weaknesses:
{", ".join(data["weaknesses"])}

--------------------------------

Suggestions:
{", ".join(data["suggestions"])}

--------------------------------

Certifications:
{", ".join(data["certifications"])}

--------------------------------

Final Verdict:

{data["final_verdict"]}
"""
                styles = getSampleStyleSheet()

                pdf = SimpleDocTemplate("Resume_Report.pdf")

                story = []

                story.append(Paragraph("Resume Analysis Report", styles["Title"]))
                story.append(
                    Paragraph(report.replace("\n", "<br/>"), styles["BodyText"])
                )

                pdf.build(story)

                document = Document()

                document.add_heading("Resume Analysis Report", level=1)

                document.add_paragraph(report)

                document.save("Resume_Report.docx")

                st.toast(
                    "🎉 Analysis Complete!",
                    icon="🎉",
                )
                col1, col2, col3 = st.columns(3)

                with col1:
                    st.download_button(
                        "📄 TXT",
                        report,
                        "Resume_Report.txt",
                        "text/plain",
                        use_container_width=True,
                    )
                with col2:
                    with open("Resume_Report.pdf", "rb") as pdf:
                        st.download_button(
                            "📕 PDF",
                            pdf,
                            "Resume_Report.pdf",
                            "application/pdf",
                            use_container_width=True,
                        )
                with col3:
                    with open("Resume_Report.docx", "rb") as doc:
                        st.download_button(
                            "📝 DOCX",
                            doc,
                            "Resume_Report.docx",
                            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True,
                        )
                st.markdown("---")
                st.caption("Built with ❤️ using Streamlit • Google Gemini • PyMuPDF")
            except json.JSONDecodeError:
                st.error("Gemini returned an invalid response. Please try again.")
            except Exception as e:
                error = str(e)
                if "RESOURCE_EXHAUSTED" in error or "429" in error:
                    st.error(
                        "🚫 Gemini API quota exceeded.\n\n"
                        "Please wait a few moments and try again, or check your Gemini API quota."
                    )
                else:
                    st.error(f"❌ {error}")
